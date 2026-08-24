"""Resume ingestion.

Extracts text from PDF/DOCX/TXT/MD and proposes CANDIDATE facts. Nothing the
parser produces is trusted: every proposed fact is stored with verified=False
and can never reach a document or an answer until a human confirms it. That is
the mechanism behind the "no fabricated qualifications" rule.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from datetime import date

from app.core.enums import FactCategory
from app.services.taxonomy import canonicalize_skill, extract_skills
from app.utils.text import normalize_ws

_EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.]+")
_PHONE_RE = re.compile(r"(?:(?:\+|00)\d{1,3}[\s.-]?)?(?:\(\d{1,4}\)[\s.-]?)?\d[\d\s.-]{6,14}\d")
_URL_RE = re.compile(r"https?://[^\s<>\"')]+", re.IGNORECASE)
_LINKEDIN_RE = re.compile(r"(?:https?://)?(?:[\w.]+\.)?linkedin\.com/in/[\w\-%]+", re.IGNORECASE)
_GITHUB_RE = re.compile(r"(?:https?://)?(?:www\.)?github\.com/[\w\-.]+", re.IGNORECASE)

_MONTHS = (
    "jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec|"
    "january|february|march|april|june|july|august|september|october|november|december"
)
_DATE_RANGE_RE = re.compile(
    rf"(?P<start>(?:{_MONTHS})?\.?\s*\d{{4}})\s*(?:-|to|\u2013|\u2014)\s*"
    rf"(?P<end>present|current|now|(?:{_MONTHS})?\.?\s*\d{{4}})",
    re.IGNORECASE,
)
_MONTH_INDEX = {
    "jan": 1,
    "feb": 2,
    "mar": 3,
    "apr": 4,
    "may": 5,
    "jun": 6,
    "jul": 7,
    "aug": 8,
    "sep": 9,
    "sept": 9,
    "oct": 10,
    "nov": 11,
    "dec": 12,
}

SECTION_PATTERNS = {
    "experience": re.compile(
        r"^\s*(work\s+)?(experience|employment|professional experience|career history)\s*:?\s*$",
        re.IGNORECASE | re.MULTILINE,
    ),
    "education": re.compile(
        r"^\s*(education|academic background|qualifications)\s*:?\s*$",
        re.IGNORECASE | re.MULTILINE,
    ),
    "skills": re.compile(
        r"^\s*(skills|technical skills|core competencies|technologies)\s*:?\s*$",
        re.IGNORECASE | re.MULTILINE,
    ),
    "certifications": re.compile(
        r"^\s*(certifications?|licenses?|accreditations?)\s*:?\s*$",
        re.IGNORECASE | re.MULTILINE,
    ),
    "projects": re.compile(r"^\s*(projects?|portfolio)\s*:?\s*$", re.IGNORECASE | re.MULTILINE),
}


@dataclass(slots=True)
class ProposedFact:
    category: str
    key: str = ""
    value: str = ""
    organization: str = ""
    title: str = ""
    start_date: date | None = None
    end_date: date | None = None
    is_current: bool = False
    highlights: list[str] = field(default_factory=list)
    evidence_url: str = ""
    #: Always False. The UI must collect an explicit human confirmation.
    verified: bool = False


@dataclass(slots=True)
class ParsedResume:
    text: str = ""
    emails: list[str] = field(default_factory=list)
    phones: list[str] = field(default_factory=list)
    links: list[str] = field(default_factory=list)
    linkedin: str = ""
    github: str = ""
    skills: list[str] = field(default_factory=list)
    sections: dict[str, str] = field(default_factory=dict)
    proposed_facts: list[ProposedFact] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    def as_dict(self) -> dict:
        return {
            "emails": self.emails,
            "phones": self.phones,
            "links": self.links,
            "linkedin": self.linkedin,
            "github": self.github,
            "skills": self.skills,
            "sections": list(self.sections),
            "proposed_fact_count": len(self.proposed_facts),
            "warnings": self.warnings,
        }


def extract_text(content: bytes, filename: str, content_type: str = "") -> tuple[str, list[str]]:
    """Return (text, warnings). Never raises on a malformed file."""
    warnings: list[str] = []
    lowered = (filename or "").lower()
    try:
        if lowered.endswith(".pdf") or "pdf" in content_type:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(content))
            pages = [(page.extract_text() or "") for page in reader.pages]
            text = "\n".join(pages)
            if not text.strip():
                warnings.append(
                    "No text layer found in the PDF. It may be a scan; upload a text-based "
                    "version or enter your facts manually. Nothing will be guessed."
                )
            return text, warnings
        if lowered.endswith(".docx") or "wordprocessingml" in content_type:
            import docx

            document = docx.Document(io.BytesIO(content))
            blocks = [p.text for p in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    blocks.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(blocks), warnings
        if lowered.endswith(".doc"):
            warnings.append("Legacy .doc is not supported. Convert to .docx or PDF.")
            return "", warnings
        return content.decode("utf-8", errors="replace"), warnings
    except Exception as exc:  # noqa: BLE001 - never fail an upload on a bad file
        warnings.append(f"Could not read {filename}: {exc}")
        return "", warnings


def _to_date(token: str, *, end: bool = False) -> date | None:
    token = (token or "").strip().lower().replace(".", "")
    if token in ("present", "current", "now"):
        return None
    match = re.search(r"(\d{4})", token)
    if not match:
        return None
    year = int(match.group(1))
    if not (1950 <= year <= date.today().year + 1):
        return None
    month = 1
    for name, index in _MONTH_INDEX.items():
        if token.startswith(name):
            month = index
            break
    else:
        month = 12 if end else 1
    return date(year, month, 1)


def split_sections(text: str) -> dict[str, str]:
    """Slice the resume at recognised headings. Unmatched text lands in 'header'."""
    marks: list[tuple[int, int, str]] = []
    for name, pattern in SECTION_PATTERNS.items():
        for match in pattern.finditer(text):
            marks.append((match.start(), match.end(), name))
    if not marks:
        return {"header": text}
    marks.sort()
    sections: dict[str, str] = {"header": text[: marks[0][0]].strip()}
    for index, (_start, heading_end, name) in enumerate(marks):
        body_end = marks[index + 1][0] if index + 1 < len(marks) else len(text)
        body = text[heading_end:body_end].strip()
        if body:
            sections[name] = (sections.get(name, "") + "\n" + body).strip()
    return {k: v for k, v in sections.items() if v}


def _bullets(block: str, limit: int = 8) -> list[str]:
    found = re.findall(r"^\s*[-*\u2022\u25cf\u00b7]\s*(.{10,300})$", block, re.MULTILINE)
    return [normalize_ws(b) for b in found][:limit]


def _propose_experience(block: str) -> list[ProposedFact]:
    facts: list[ProposedFact] = []
    lines = [line.rstrip() for line in block.splitlines()]
    current: ProposedFact | None = None
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        match = _DATE_RANGE_RE.search(stripped)
        if match:
            headline = normalize_ws(_DATE_RANGE_RE.sub("", stripped).strip(" |,-\u2013"))
            title, _, organization = (
                headline.partition(" at ") if " at " in headline else headline.partition(",")
            )
            end_token = match.group("end").lower()
            current = ProposedFact(
                category=FactCategory.EMPLOYMENT.value,
                title=normalize_ws(title)[:200],
                organization=normalize_ws(organization)[:200],
                value=headline[:500],
                start_date=_to_date(match.group("start")),
                end_date=_to_date(end_token, end=True),
                is_current=end_token in ("present", "current", "now"),
            )
            facts.append(current)
        elif current is not None and re.match(r"^\s*[-*\u2022\u25cf\u00b7]", line):
            bullet = normalize_ws(re.sub(r"^\s*[-*\u2022\u25cf\u00b7]\s*", "", line))
            if len(bullet) >= 10:
                current.highlights.append(bullet[:300])
    return facts


def _propose_education(block: str) -> list[ProposedFact]:
    facts: list[ProposedFact] = []
    for line in block.splitlines():
        stripped = normalize_ws(line)
        if len(stripped) < 8:
            continue
        if not re.search(
            r"\b(b\.?s|b\.?a|b\.?tech|m\.?s|m\.?a|m\.?tech|mba|ph\.?d|bachelor|master|"
            r"doctor|diploma|degree|university|college|institute)\b",
            stripped,
            re.IGNORECASE,
        ):
            continue
        match = _DATE_RANGE_RE.search(stripped)
        facts.append(
            ProposedFact(
                category=FactCategory.EDUCATION.value,
                value=stripped[:500],
                organization=normalize_ws(
                    next(
                        (
                            part
                            for part in re.split(r"[,|]", stripped)
                            if re.search(r"university|college|institute|school", part, re.I)
                        ),
                        "",
                    )
                )[:200],
                start_date=_to_date(match.group("start")) if match else None,
                end_date=_to_date(match.group("end"), end=True) if match else None,
            )
        )
    return facts[:10]


def parse_resume(content: bytes, filename: str, content_type: str = "") -> ParsedResume:
    """Extract text and propose UNVERIFIED facts for human confirmation."""
    text, warnings = extract_text(content, filename, content_type)
    result = ParsedResume(text=text, warnings=warnings)
    if not text.strip():
        return result

    result.emails = list(dict.fromkeys(_EMAIL_RE.findall(text)))[:5]
    result.phones = [
        normalize_ws(p)
        for p in dict.fromkeys(_PHONE_RE.findall(text))
        if len(re.sub(r"\D", "", p)) >= 8
    ][:3]
    result.links = list(dict.fromkeys(_URL_RE.findall(text)))[:20]
    linkedin = _LINKEDIN_RE.search(text)
    github = _GITHUB_RE.search(text)
    result.linkedin = linkedin.group(0) if linkedin else ""
    result.github = github.group(0) if github else ""
    result.sections = split_sections(text)
    result.skills = extract_skills(text)

    proposed: list[ProposedFact] = []
    proposed += _propose_experience(result.sections.get("experience", ""))
    proposed += _propose_education(result.sections.get("education", ""))
    for skill in result.skills:
        proposed.append(
            ProposedFact(
                category=FactCategory.SKILL.value,
                key=canonicalize_skill(skill),
                value=skill,
            )
        )
    for block in _bullets(result.sections.get("certifications", "")):
        proposed.append(ProposedFact(category=FactCategory.CERTIFICATION.value, value=block))
    for link in (result.linkedin, result.github):
        if link:
            proposed.append(
                ProposedFact(category=FactCategory.LINK.value, value=link, evidence_url=link)
            )

    result.proposed_facts = proposed
    if not proposed:
        result.warnings.append(
            "No facts could be proposed from this file. Add them manually; the agent will "
            "not invent anything."
        )
    result.warnings.append(
        f"{len(proposed)} facts proposed and stored UNVERIFIED. Verify each one before it "
        "can be used in an application."
    )
    return result
